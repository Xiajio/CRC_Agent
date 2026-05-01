from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
ASSET_DIR = DOCS_DIR / "_architecture_docx_assets"
OUTPUT = DOCS_DIR / "current-architecture-map.docx"

ACCENT = "1F4E79"
ACCENT_DARK = "173B5F"
ACCENT_LIGHT = "D9EAF7"
MUTED = "5B6770"
TABLE_HEADER = "EAF3FA"
BOX_FILL = (236, 246, 252)
BOX_OUTLINE = (45, 95, 135)
ARROW = (60, 70, 80)


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def find_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\segoeuib.ttf" if bold else r"C:\Windows\Fonts\segoeui.ttf",
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def wrap_text(text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = word if not current else f"{current} {word}"
        bbox = font.getbbox(candidate)
        if bbox[2] - bbox[0] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines or [text]


def box_center(box: tuple[int, int, int, int]) -> tuple[int, int]:
    x, y, w, h = box
    return x + w // 2, y + h // 2


def edge_points(
    source: tuple[int, int, int, int],
    target: tuple[int, int, int, int],
) -> tuple[tuple[int, int], tuple[int, int]]:
    sx, sy = box_center(source)
    tx, ty = box_center(target)
    x, y, w, h = source
    tx0, ty0, tw, th = target
    if abs(tx - sx) >= abs(ty - sy):
        start = (x + w, sy) if tx >= sx else (x, sy)
        end = (tx0, ty) if tx >= sx else (tx0 + tw, ty)
    else:
        start = (sx, y + h) if ty >= sy else (sx, y)
        end = (tx, ty0) if ty >= sy else (tx, ty0 + th)
    return start, end


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    draw.line([start, end], fill=ARROW, width=3)
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex >= sx else -1
        points = [(ex, ey), (ex - direction * 12, ey - 7), (ex - direction * 12, ey + 7)]
    else:
        direction = 1 if ey >= sy else -1
        points = [(ex, ey), (ex - 7, ey - direction * 12), (ex + 7, ey - direction * 12)]
    draw.polygon(points, fill=ARROW)


def draw_flowchart(
    name: str,
    size: tuple[int, int],
    title: str,
    nodes: dict[str, tuple[int, int, int, int, str]],
    edges: Iterable[tuple[str, str]],
) -> Path:
    image = Image.new("RGB", size, "white")
    draw = ImageDraw.Draw(image)
    title_font = find_font(30, bold=True)
    node_font = find_font(20)
    small_font = find_font(17)

    draw.text((42, 28), title, fill=(20, 60, 90), font=title_font)
    draw.line([(42, 70), (size[0] - 42, 70)], fill=(210, 225, 235), width=2)

    node_boxes = {key: value[:4] for key, value in nodes.items()}
    for source, target in edges:
        start, end = edge_points(node_boxes[source], node_boxes[target])
        draw_arrow(draw, start, end)

    for key, (x, y, w, h, text) in nodes.items():
        fill = (226, 239, 248) if key.startswith("g") else BOX_FILL
        draw.rounded_rectangle(
            (x, y, x + w, y + h),
            radius=16,
            fill=fill,
            outline=BOX_OUTLINE,
            width=3,
        )
        lines = wrap_text(text, node_font, w - 24)
        if len(lines) > 3:
            lines = wrap_text(text, small_font, w - 24)
            font = small_font
        else:
            font = node_font
        line_height = max(font.getbbox("Ag")[3] - font.getbbox("Ag")[1] + 5, 22)
        total_height = len(lines) * line_height
        ty = y + max(10, (h - total_height) // 2)
        for line in lines:
            bbox = font.getbbox(line)
            tw = bbox[2] - bbox[0]
            draw.text((x + (w - tw) // 2, ty), line, fill=(23, 42, 58), font=font)
            ty += line_height

    output = ASSET_DIR / f"{name}.png"
    image.save(output, quality=95)
    return output


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9.5)
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if widths:
        for idx, width in enumerate(widths):
            table.columns[idx].width = Cm(width)
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], TABLE_HEADER)
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    document.add_paragraph()


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.color.rgb = RGBColor.from_string(ACCENT if level == 1 else ACCENT_DARK)


def add_body(document: Document, text: str, style: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10)


def add_diagram(document: Document, path: Path, caption: str, width_inches: float = 6.7) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    caption_p = document.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_p.add_run(caption)
    caption_run.font.name = "Microsoft YaHei"
    caption_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    caption_run.font.size = Pt(9)
    caption_run.font.color.rgb = RGBColor.from_string(MUTED)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True


def build_diagrams() -> dict[str, Path]:
    ensure_dir(ASSET_DIR)
    diagrams: dict[str, Path] = {}

    diagrams["runtime"] = draw_flowchart(
        "runtime_initialization",
        (1280, 760),
        "后端运行时装配",
        {
            "a": (55, 120, 230, 72, "应用创建入口"),
            "b": (360, 120, 210, 72, "接口服务实例"),
            "c": (665, 120, 220, 72, "启动生命周期"),
            "d": (975, 120, 230, 72, "运行时容器"),
            "e": (70, 300, 205, 78, "配置与鉴权"),
            "f": (320, 300, 205, 78, "会话内存仓库"),
            "g1": (570, 300, 205, 78, "患者登记库"),
            "g2": (820, 300, 205, 78, "上传资产目录"),
            "h": (360, 500, 210, 78, "患者图服务"),
            "i": (665, 500, 210, 78, "医生图服务"),
            "j": (975, 500, 230, 78, "场景分流器"),
        },
        [
            ("a", "b"),
            ("b", "c"),
            ("c", "d"),
            ("c", "e"),
            ("c", "f"),
            ("c", "g1"),
            ("c", "g2"),
            ("f", "h"),
            ("f", "i"),
            ("h", "j"),
            ("i", "j"),
        ],
    )

    diagrams["stream"] = draw_flowchart(
        "graph_streaming",
        (1280, 610),
        "单轮对话流式执行",
        {
            "a": (50, 155, 170, 78, "发送流式消息"),
            "b": (270, 155, 170, 78, "场景选择"),
            "c": (490, 155, 170, 78, "图服务执行"),
            "d": (710, 155, 170, 78, "运行锁与上下文"),
            "e": (930, 155, 170, 78, "图引擎流输出"),
            "f": (490, 350, 170, 78, "事件标准化"),
            "g": (710, 350, 170, 78, "编码为 SSE"),
            "h": (930, 350, 170, 78, "流式响应返回"),
        },
        [
            ("a", "b"),
            ("b", "c"),
            ("c", "d"),
            ("d", "e"),
            ("e", "f"),
            ("f", "g"),
            ("g", "h"),
        ],
    )

    diagrams["ingest"] = draw_flowchart(
        "rag_ingest",
        (1280, 680),
        "知识库离线构建",
        {
            "a": (55, 130, 190, 76, "指南源文件"),
            "b": (295, 130, 190, 76, "文档解析"),
            "c": (535, 130, 190, 76, "元数据抽取"),
            "d": (775, 130, 190, 76, "文本切分"),
            "e": (1015, 130, 190, 76, "片段元数据"),
            "f": (295, 360, 190, 76, "向量化"),
            "g1": (535, 360, 190, 76, "向量库写入"),
            "g2": (775, 360, 190, 76, "词法索引构建"),
            "h1": (535, 520, 190, 76, "向量库落盘"),
            "h2": (775, 520, 190, 76, "词法索引落盘"),
        },
        [
            ("a", "b"),
            ("b", "c"),
            ("c", "d"),
            ("d", "e"),
            ("e", "f"),
            ("f", "g1"),
            ("e", "g2"),
            ("g1", "h1"),
            ("g2", "h2"),
        ],
    )

    diagrams["retrieval"] = draw_flowchart(
        "rag_retrieval",
        (1280, 680),
        "RAG 混合检索",
        {
            "a": (55, 130, 190, 76, "知识检索工具"),
            "b": (295, 130, 190, 76, "统一检索入口"),
            "c": (535, 130, 190, 76, "检索器管理"),
            "d": (775, 130, 190, 76, "混合检索器"),
            "e": (295, 350, 190, 76, "向量召回"),
            "f": (535, 350, 190, 76, "词法召回"),
            "g": (775, 350, 190, 76, "结果融合"),
            "h": (1015, 350, 190, 76, "可选重排"),
            "i": (775, 520, 190, 76, "格式化证据"),
        },
        [
            ("a", "b"),
            ("b", "c"),
            ("c", "d"),
            ("d", "e"),
            ("d", "f"),
            ("e", "g"),
            ("f", "g"),
            ("g", "h"),
            ("h", "i"),
        ],
    )

    diagrams["frontend"] = draw_flowchart(
        "frontend_sse",
        (1280, 630),
        "前端 API 与 SSE 消费",
        {
            "a": (55, 145, 190, 76, "页面提交消息"),
            "b": (295, 145, 190, 76, "API 客户端"),
            "c": (535, 145, 190, 76, "POST 请求"),
            "d": (775, 145, 190, 76, "读取响应流"),
            "e": (1015, 145, 190, 76, "解析 SSE 数据"),
            "f": (295, 365, 190, 76, "调试旁路"),
            "g": (535, 365, 190, 76, "事件回调"),
            "h": (775, 365, 190, 76, "状态归并"),
            "i": (1015, 365, 190, 76, "界面状态"),
        },
        [
            ("a", "b"),
            ("b", "c"),
            ("c", "d"),
            ("d", "e"),
            ("e", "f"),
            ("e", "g"),
            ("g", "h"),
            ("h", "i"),
        ],
    )

    diagrams["scenes"] = draw_flowchart(
        "scene_boundaries",
        (1280, 700),
        "患者与医生场景边界",
        {
            "a": (70, 145, 210, 78, "患者会话"),
            "b": (370, 145, 210, 78, "草稿患者身份"),
            "c": (670, 145, 210, 78, "身份与上传资料"),
            "d": (970, 145, 210, 78, "患者图运行"),
            "e": (70, 420, 210, 78, "医生会话"),
            "f": (370, 420, 210, 78, "绑定登记患者"),
            "g": (670, 420, 210, 78, "注入患者摘要"),
            "h": (970, 420, 210, 78, "医生图运行"),
        },
        [
            ("a", "b"),
            ("b", "c"),
            ("c", "d"),
            ("e", "f"),
            ("f", "g"),
            ("g", "h"),
        ],
    )

    return diagrams


def build_document() -> None:
    diagrams = build_diagrams()
    document = Document()
    configure_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("LangG 现状架构图谱说明")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(ACCENT_DARK)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    srun = subtitle.add_run("面向项目内部的当前态说明：代码入口、数据流、落盘边界与 SSE 消费")
    srun.font.name = "Microsoft YaHei"
    srun._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    srun.font.size = Pt(12)
    srun.font.color.rgb = RGBColor.from_string(MUTED)

    add_table(
        document,
        ["范围", "说明"],
        [
            ["后端入口", "backend/app.py、backend/api/routes/*、backend/api/services/graph_service.py"],
            ["RAG 入口", "src/rag/retriever.py、src/tools/rag_tools.py；构建链路参考 src/rag/ingest.py 与 bm25_index.py"],
            ["前端入口", "frontend/src/app/api/client.ts、frontend/src/pages/workspace-page.tsx、stream-reducer.ts"],
            ["本文边界", "描述当前实现，不写目标架构；重点放在运行入口、数据流、持久化和关键边界。"],
        ],
        widths=[4.0, 12.5],
    )

    add_body(
        document,
        "结论摘要：当前系统是一个 FastAPI BFF 包装两套 LangGraph 运行图。session 元数据在进程内，患者登记和上传资产落盘在 runtime/，RAG 知识库通过离线 ingest 构建到 Chroma 与 BM25，前端通过 POST SSE 读取 graph 事件并用 reducer 落到页面状态。",
    )

    document.add_page_break()

    add_heading(document, "1. 总体运行入口", 1)
    add_diagram(document, diagrams["runtime"], "图 1：后端启动后装配运行时容器、图服务与场景分流器。")
    add_body(
        document,
        "backend/app.py:create_app() 创建 FastAPI 实例，挂载鉴权、CORS 与业务路由。lifespan 中创建 runtime_root、assets_root、PatientRegistryService、InMemorySessionStore、patient graph、doctor graph 以及 SceneGraphRouter，并集中放入 app.state.runtime。",
    )
    add_bullets(
        document,
        [
            "runner_mode 默认 real；GRAPH_RUNNER_MODE=fixture 时 patient/doctor 都走 FixtureGraphRunner。",
            "RAG_WARMUP 默认开启，但只 warmup retriever，不自动构建知识库。",
            "session_store 是内存态；patient_registry_service 指向 runtime/patient_registry.db；上传资产写入 runtime/assets/。",
        ],
    )

    add_heading(document, "2. API 路由职责", 1)
    add_table(
        document,
        ["Router", "Prefix", "当前职责"],
        [
            ["sessions.py", "/api/sessions", "创建、查询、重置 session；读取消息历史；doctor 绑定患者；patient 保存身份。"],
            ["chat.py", "/api/sessions", "POST /{session_id}/messages/stream，以 text/event-stream 运行一轮 graph。"],
            ["database.py", "/api/database", "医生侧历史病例库统计、搜索、详情、upsert 与 query intent。"],
            ["patient_registry.py", "/api/patient-registry", "患者登记库 recent/search/detail/records/alerts/delete/clear。"],
            ["uploads.py", "/api", "患者侧上传文件，转换 medical_card，写入 session context 和患者登记库。"],
            ["assets.py", "/api", "按 asset_id 读取上传资产内容，设置安全的 Content-Disposition。"],
        ],
        widths=[3.2, 4.0, 9.3],
    )
    add_body(
        document,
        "路由层不直接运行 LangGraph。chat.py 通过 request.app.state.runtime.scene_router.for_session(session_id) 选择 GraphService，然后将其 async iterator 包装成 StreamingResponse。",
    )

    add_heading(document, "3. patient/doctor 运行时调度", 1)
    add_diagram(document, diagrams["scenes"], "图 2：patient 与 doctor 场景的上下文来源和注入边界。")
    add_heading(document, "3.1 session 与场景分流", 2)
    add_body(
        document,
        "sessions.py:create_session() 只接受 patient 或 doctor。patient session 创建时会在 runtime/patient_registry.db 中创建草稿患者，并把 patient_id 写回 SessionMeta；doctor session 默认只创建内存 session。",
    )
    add_body(
        document,
        "SceneGraphRouter.for_session(session_id) 是 graph 分流点：meta.scene == patient 时使用 PatientGraphService，其余使用 DoctorGraphService。",
    )
    add_heading(document, "3.2 单轮 graph 执行", 2)
    add_diagram(document, diagrams["stream"], "图 3：单轮消息从 API 入口进入图执行，再被编码成 SSE 返回。")
    add_bullets(
        document,
        [
            "GraphService.stream_turn() 先读取 SessionMeta，再通过 try_acquire_run_lock 防止同 session 并发运行。",
            "payload_builder.build_graph_payload() 合成 messages、patient_profile、findings、roadmap、medical_card、summary_memory 等 graph 输入。",
            "compiled_graph.astream() 使用 thread_id 作为 LangGraph configurable thread_id。",
            "模型 token callback 转为 message.delta；node output 通过 normalize_tick() 转为 status.node、message.done、card.upsert、references.append 等事件。",
            "成功后 bump snapshot_version 并发送 done；失败时恢复 pending_context_messages，发送 error 和 done。",
        ],
    )
    add_heading(document, "3.3 doctor 特有上下文", 2)
    add_body(
        document,
        "DoctorGraphService._prepare_session_meta() 在医生 session 绑定 patient_id 且尚未注入过该患者时，从 PatientRegistryService 读取患者摘要与 alerts，组合成 HumanMessage 放入 pending_context_messages。下一轮 graph payload 会先带上这条上下文消息。",
    )
    add_heading(document, "3.4 patient 特有上下文", 2)
    add_body(
        document,
        "PatientGraphService 不启用 context finalizer，也不主动读取 patient registry 摘要。患者侧上下文主要来自 session 创建的 draft patient_id、identity 接口和 uploads 接口。上传后 upload_service 会把 medical_card 写入 session context_state，并在需要时写入 patient_registry.db。",
    )

    document.add_page_break()

    add_heading(document, "4. 知识库构建与落盘", 1)
    add_diagram(document, diagrams["ingest"], "图 4：离线构建从指南源文件生成向量库与词法索引。")
    add_body(
        document,
        "知识库构建入口是 src/rag/ingest.py，不在 backend startup 中自动执行。常用命令是 python -m src.rag.ingest；需要重建时使用 --reset。",
    )
    add_table(
        document,
        ["对象", "位置 / 配置", "说明"],
        [
            ["原始文档", "data/guidelines/*.pdf|*.md|*.txt", "ingest 扫描的知识库源文件。"],
            ["Chroma collection", "clinical_guidelines", "通过 Chroma.add_documents(ids=chunk_id) 写入。"],
            ["Chroma 目录", "RAG_PERSIST_DIR 或 settings.rag.persist_dir；默认 chroma_db/", "向量库持久化位置。"],
            ["BM25 目录", "RAG_BM25_INDEX_PATH 或 settings.rag.bm25_index_path；默认 bm25_index/", "词法索引持久化位置。"],
            ["BM25 文件", "bm25_data.pkl.gz、bm25_data.meta.json", "pickle+gzip 数据与可读 metadata。"],
            ["Embedding 后端", "RAG_EMBEDDING_BACKEND=api/local", "api 走 EMBEDDING_API_BASE/KEY；local 走 Hugging Face 模型。"],
        ],
        widths=[3.6, 6.3, 6.6],
    )
    add_body(
        document,
        "ingest 会解析文档、抽取文档级 metadata、切分 chunk、补充 source/chunk_id/page/section 等元数据，可选地加入 contextual prefix 和 hypothetical questions，然后同一批 chunks 同步写入 Chroma 与 BM25。",
    )

    add_heading(document, "5. RAG 检索链路", 1)
    add_diagram(document, diagrams["retrieval"], "图 5：RAG 工具入口经过混合召回、融合和可选重排形成证据。")
    add_heading(document, "5.1 工具注册", 2)
    add_body(
        document,
        "src/tools/rag_tools.py 将检索能力封装为 LangChain BaseTool。get_enhanced_rag_tools() 当前默认返回 ClinicalGuidelineSearchTool、TreatmentSearchTool、StagingSearchTool、DrugInfoSearchTool、GuidelineStructureTool 和 GuidelineReaderTool。",
    )
    add_heading(document, "5.2 从 graph 到 retriever", 2)
    add_body(
        document,
        "src/nodes/knowledge_nodes.py:node_knowledge_retrieval() 从工具列表里找 search_clinical_guidelines 作为 local RAG。plan-driven 模式按 current_step.tool_needed 调用 TOC、chapter、search 或 treatment 工具；普通模式下，如果问题需要患者上下文，会先调用 local_rag_tool.invoke({query, top_k})。",
    )
    add_heading(document, "5.3 retriever 内部", 2)
    add_bullets(
        document,
        [
            "_get_vectorstore() 用 lru_cache 缓存 Chroma collection。",
            "_GlobalRetrieverManager 单例缓存 SimpleRetriever、vectorstore 和 reranker。",
            "SimpleRetriever.retrieve() 默认执行向量检索 + BM25 检索 + rank fusion，可选 rerank。",
            "metadata filter 支持 $eq、$in、$ne、$and、$or。",
            "工具输出先被格式化为带 [[Source:file|Page:n]] 和 retrieved_metadata 的文本，再由 graph synthesis 生成用户可见回答和 references.append。",
        ],
    )

    document.add_page_break()

    add_heading(document, "6. 前端 API 调用与 SSE 消费", 1)
    add_diagram(document, diagrams["frontend"], "图 6：前端页面通过 API 客户端消费 POST SSE 流式事件。")
    add_body(
        document,
        "frontend/src/app/api/client.ts:createApiClient() 是前端请求集中入口。普通 JSON API 使用 parseJsonResponse；流式消息通过 streamTurn() 调用 streamJsonEvents()。",
    )
    add_body(
        document,
        "当前没有使用浏览器 EventSource，因为服务端接口是 POST JSON body。frontend/src/app/api/stream.ts 使用 fetch + response.body.getReader()，按空行分割 SSE block，只解析 data: 行为 StreamEvent。",
    )
    add_table(
        document,
        ["事件", "前端 reducer 行为"],
        [
            ["status.node", "更新 statusNode，并推进 roadmap。"],
            ["message.delta", "创建或追加流式 AI 消息。"],
            ["message.done", "写入最终 AI 消息、thinking、inline cards。"],
            ["card.upsert", "写入 cards，并尽量挂到最新 AI 消息。"],
            ["references.append", "去重追加 references。"],
            ["context.maintenance", "更新 context maintenance 状态，doctor 侧会触发轮询 getSession。"],
            ["error", "写入 lastError，并把当前 plan step 标记 blocked。"],
            ["done", "更新 threadId、snapshotVersion，清理 active run/status。"],
        ],
        widths=[4.2, 12.3],
    )
    add_body(
        document,
        "WorkspacePage 通过 useSceneSessions() 同时维护 patient 与 doctor 两个 session id，并写入 localStorage。后端重启导致内存 session 丢失时，前端收到 404 会清理 stale id 并重新创建 session。",
    )

    add_heading(document, "7. 关键数据与边界", 1)
    add_table(
        document,
        ["数据", "位置", "生命周期 / 边界"],
        [
            ["session id/thread id/active run/context_state", "InMemorySessionStore", "进程内；后端重启丢失。"],
            ["LangGraph checkpoint", "get_checkpointer(settings.checkpoint)", "取决于 checkpoint 配置。"],
            ["患者登记库", "runtime/patient_registry.db", "SQLite 落盘。"],
            ["上传资产", "runtime/assets/", "文件落盘。"],
            ["Chroma 向量库", "RAG_PERSIST_DIR 或 chroma_db/", "离线构建后运行时只读。"],
            ["BM25 词法索引", "RAG_BM25_INDEX_PATH 或 bm25_index/", "离线构建后运行时只读。"],
            ["前端 session id", "localStorage", "浏览器本地；后端 session 丢失时重建。"],
        ],
        widths=[5.6, 5.2, 5.7],
    )
    add_bullets(
        document,
        [
            "每个 session 只有一个 active run；chat stream 与 upload 都会使用 session run lock。",
            "patient scene 偏自述、身份、上传与轻量 graph；doctor scene 可绑定 registry patient，并在下一轮运行前注入患者摘要。",
            "RAG ingest 是离线动作；runtime 只初始化和缓存 retriever。",
            "citations/references 的结构化传递依赖 RAG tool 输出和 graph 节点提取；前端只消费 references.append。",
        ],
    )

    section = document.add_section(WD_SECTION.CONTINUOUS)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    document.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
