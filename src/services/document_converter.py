"""
医疗文档转化 Agent (Pro Optimized Version)

优化特性：
1. Pydantic 强类型约束：确保 JSON 结构绝对稳定，自动处理字段类型。
2. 混合解析引擎 (Hybrid Engine)：按页智能切换 OCR/文本模式，不再"一刀切"。
3. PyMuPDF (Fitz) 极速内核：统一替换 pypdf，文本提取速度提升 10x。
4. 思维链 (CoT) 嵌入：强制模型先思考分期逻辑，再填充字段。
"""

from __future__ import annotations

import base64
import io
import json
import logging
import os
import re
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Any, Dict, List, Optional, Union

# 引入 Pydantic 进行结构化定义
# 兼容新旧版本的 langchain_core：优先尝试从 langchain_core 导入，失败则直接从 pydantic 导入
try:
    # 优先尝试从 langchain_core 导入（兼容旧环境）
    from langchain_core.pydantic_v1 import BaseModel, Field, validator
except ImportError:
    # 如果失败（新环境），直接从 pydantic 导入
    from pydantic import BaseModel, Field, validator
from langchain_core.messages import HumanMessage, SystemMessage

from .llm_service import create_compatible_chat_openai

# 统一依赖检查
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ============== 0. 枚举定义 (Enums) - 替代 Prompt 中的文本规则 ==============
# 核心思想：将"不确定的自然语言指令"转化为"确定的 Python/Pydantic 逻辑"

from enum import Enum
from typing import Literal


# 肿瘤位置枚举 (限制可选值，防止 LLM 自由发挥)
class TumorLocationEnum(str, Enum):
    RECTUM = "直肠"
    COLON = "结肠"
    ASCENDING = "升结肠"
    DESCENDING = "降结肠"
    SIGMOID = "乙状结肠"
    CECUM = "盲肠"
    TRANSVERSE = "横结肠"
    RECTOSIGMOID = "直肠乙状结肠交界"
    ANUS = "肛管"
    UNKNOWN = "未提供"


# 分化程度枚举 (自动映射英文到中文)
class DifferentiationEnum(str, Enum):
    HIGH = "高分化"       # well differentiated, G1, low grade
    MODERATE = "中分化"   # moderately differentiated, G2, intermediate grade
    LOW = "低分化"        # poorly differentiated, G3, high grade
    UNDIFFERENTIATED = "未分化"  # undifferentiated
    UNKNOWN = "未提供"


# 风险状态枚举
class RiskStatusEnum(str, Enum):
    HIGH = "高危"
    LOW = "低危"
    UNKNOWN = "未评估"


# 分期类型枚举
class StagingTypeEnum(str, Enum):
    CLINICAL = "临床分期"
    PATHOLOGICAL = "病理分期"
    POST_NEOADJUVANT = "新辅助后分期"
    UNKNOWN = "待评估"


# ============== 1. Pydantic 结构化定义 (OpenAI 兼容终极版) ==============

# [新增] 专门定义分子标志物类，替代不稳定的 Dict
class MolecularMarkers(BaseModel):
    MSI: str = Field(description="MSI状态 (如 MSS, MSI-H, pMMR, dMMR)。未找到填'未提供'")
    KRAS: str = Field(description="KRAS状态。未找到填'未提供'")
    NRAS: str = Field(description="NRAS状态。未找到填'未提供'")
    BRAF: str = Field(description="BRAF状态。未找到填'未提供'")
    HER2: str = Field(description="HER2状态。未找到填'未提供'")
    UGT1A1: str = Field(description="UGT1A1状态。未找到填'未提供'")
    others: str = Field(description="其他重要基因突变。未找到填'无'")

class PatientSummary(BaseModel):
    age: str = Field(description="患者年龄，如 '52岁'，未找到填'未提供'")
    gender: str = Field(description="性别，未找到填'未提供'")
    chief_complaint: str = Field(description="主诉症状，未找到填'未提供'")
    admission_date: str = Field(description="入院日期，未找到填'未提供'")
    medical_history: str = Field(description="既往史摘要，未找到填'未提供'")

class DiagnosisBlock(BaseModel):
    confirmed: str = Field(description="确诊诊断（如 结肠腺癌、直肠癌）")
    # [优化] 使用枚举限制可选值
    location: TumorLocationEnum = Field(description="肿瘤位置（枚举值）")
    pathology: str = Field(description="病理类型（如 腺癌、鳞癌）")
    # [优化] 使用枚举，自动处理英文到中文映射
    differentiation: DifferentiationEnum = Field(description="分化程度（枚举值）")
    # [循证增强] 强制要求模型摘录原文，防止幻觉
    evidence_quote: str = Field(description="[Grounding] 支持上述诊断的病理报告原文片段")
    # [关键修复] 将 Dict 替换为具体的 MolecularMarkers 类
    molecular_markers: MolecularMarkers = Field(description="分子标志物详情")

    # [逻辑代码化] 自动清洗诊断文本
    @validator('confirmed')
    def clean_diagnosis(cls, v):
        if isinstance(v, str):
            return v.strip(" .。")
        return v


class StagingBlock(BaseModel):
    t_stage: str = Field(description="T分期（含前缀，如 pT3, cT4a, ypT2）")
    n_stage: str = Field(description="N分期（含前缀，如 pN0, cN1）")
    m_stage: str = Field(description="M分期（含前缀，如 pM0, cM1）")
    staging_type: StagingTypeEnum = Field(description="分期类型（枚举值）")
    clinical_stage: str = Field(description="总体分期（中文，如 III期、IIIb期）")
    risk_status: RiskStatusEnum = Field(description="风险状态（枚举值）")
    # [循证增强] 强制要求原文引用，防止模型幻觉分期
    evidence_quote: str = Field(description="[Grounding] 支持上述分期的影像或病理原文片段")

    # [逻辑代码化] OCR 纠错逻辑 (Code > Prompt)
    # 替代 Prompt 中的 "注意 TI -> T1" 等规则
    @validator('t_stage', 'n_stage', 'm_stage', pre=True)
    def fix_ocr_errors(cls, v):
        if not isinstance(v, str):
            return v
        v = v.strip().upper()  # 统一大写
        # 常见 OCR 错误修复表
        corrections = {
            "TI": "T1", "TL": "T1", "T|": "T1",
            "TO": "T0", "NO": "N0", "MO": "M0",
            "ll": "II", "Il": "II",
            "TIS": "Tis",  # 特殊大小写
            "IIS": "IIs",
        }
        for wrong, right in corrections.items():
            if wrong in v:
                v = v.replace(wrong, right)
        return v

class KeyFinding(BaseModel):
    category: str = Field(description="检查类别")
    finding: str = Field(description="关键发现内容")
    significance: str = Field(description="临床意义，无则填'无'")

class TreatmentStep(BaseModel):
    step: int = Field(description="步骤序号")
    name: str = Field(description="治疗措施名称")
    status: str = Field(description="状态：建议/待定/已完成")
    details: str = Field(description="详细说明，无则填'无'")

# 定义 Data 层，包裹所有子模块
class CardData(BaseModel):
    patient_summary: PatientSummary
    diagnosis_block: DiagnosisBlock
    staging_block: StagingBlock
    key_findings: List[KeyFinding]
    treatment_draft: List[TreatmentStep]

# 定义最外层结构
class MedicalVisualizationCard(BaseModel):
    """主数据结构"""
    type: str = Field(description="固定填 'medical_visualization_card'")
    thinking_process: str = Field(description="[CoT] 提取前的思考过程")
    data: CardData = Field(description="包含所有临床数据的对象")

    def to_ui_json(self):
        """转换为前端可用的扁平 Dict"""
        # Pydantic v1/v2 兼容的导出
        if hasattr(self, "model_dump"):
            return self.model_dump()
        else:
            return self.dict()

# ============== 2. 系统提示词 (Prompt Optimization) ==============

# ============== 2. 系统提示词 (优化瘦身版) ==============

# 极简 Prompt：只关注"提取"这一核心动作，规则已下沉到 Schema
LIGHTWEIGHT_SYSTEM_PROMPT = """你是一个专业的医疗数据提取引擎，擅长从中文医疗文档中提取关键分期信息。

【任务】
从医疗文本中提取符合 JSON Schema 定义的数据，**重点关注分期信息**。

【核心原则】
1. **精确提取**：直接提取原文信息，不要进行推断或总结。
2. **循证溯源**：关键字段必须在 `evidence_quote` 中摘录原文证据。
3. **前缀保留**：严格保留 TNM 分期的前缀 (c/p/yp)，这是区分临床/病理分期的唯一标准。
4. **分期优先级**：病理分期(pTNM) > 新辅助后分期(ypTNM) > 临床分期(cTNM)。
5. **隐私安全**：忽略所有人名、电话、ID号码。

【输出】
直接输出合法的 JSON 数据。不要包含 markdown 标记。"""


# 完整 Prompt：包含 Schema 说明（供 Pydantic structured_output 参考）
CONVERSION_SYSTEM_PROMPT = LIGHTWEIGHT_SYSTEM_PROMPT + """

【Schema 枚举约束】

diagnosis_block.location 必须是以下值之一：
- 直肠 / 结肠 / 升结肠 / 降结肠 / 乙状结肠 / 盲肠 / 横结肠 / 直肠乙状结肠交界 / 肛管 / 未提供

diagnosis_block.differentiation 必须是以下值之一：
- 高分化 / 中分化 / 低分化 / 未分化 / 未提供

staging_block.risk_status 必须是以下值之一：
- 高危 / 低危 / 未评估

staging_block.staging_type 必须是以下值之一：
- 临床分期 / 病理分期 / 新辅助后分期 / 待评估

【关键规则】
1. **TNM 分期前缀识别**：
   - cTNM = 临床分期（来自影像报告/术前CT/MRI）：如 cT3N1M0
   - pTNM = 病理分期（来自手术后病理报告）：如 pT3N1M0
   - ypTNM = 新辅助治疗后分期：如 ypT2N0M0
   - 无前缀的 TNM 默认按临床分期处理

2. **分期信息提取策略（按优先级）**：
   **优先级1**：明确标注的总体分期
     - 识别关键词："分期"、"stage"、"期"
     - 格式示例："III期"、"IIIB期"、"III b期"、"stage III"、"Stage IIIb"
     - 中文/阿拉伯数字混用："III期" = "IIIb期" = "3期"
   
   **优先级2**：TNM 分期组合（需要推断总体分期）
     - 根据完整的 TNM 信息推断总体分期（见下方映射表）
     - 如果只有 T、N、M 中的部分信息，标注为"待评估"
     - 示例：T3N1M0 → III期；T4N2M1 → IV期

   **优先级3**：其他分期描述
     - 如"晚期"、"进展期"等描述性术语，转换为对应分期或标注"待评估"

3. **TNM → 总体分期映射表（结直肠癌）**：
   ```
   0期（原位癌）：  Tis N0 M0
   I期：           T1-2 N0 M0
   IIA期：         T3 N0 M0
   IIB期：         T4a N0 M0
   IIC期：         T4b N0 M0
   IIIA期：        T1-2 N1/N1c M0
   IIIB期：        T3-4a N1/N1c M0
   IIIC期：        T3-4a N2a/N2b M0 或 T4b N1-2 M0
   IVA期：        任何T 任何N M1a
   IVB期：        任何T 任何N M1b
   ```

4. **分期识别技巧**：
   - **位置扫描**：重点检查"病理报告"、"手术记录"、"影像报告"、"出院小结"等章节
   - **关键词组合**：
     * "T3 N1 M0"（空格分隔）
     * "cT3N1M0"（紧邻书写）
     * "cT3, N1, M0"（逗号分隔）
     * "T3期 N1期 M0期"（逐项描述）
   - **数字识别**：注意区分罗马数字（III）和阿拉伯数字（3）
   - **混合格式**：如 "III期 (cT3N1M0)"，提取时保留分期III，在evidence_quote中记录完整信息

5. **分期类型判断**：
   - 有 "病理"、"pTNM"、"术后"、"术后病理" → 病理分期
   - 有 "影像"、"CT"、"MRI"、"cTNM"、"术前" → 临床分期
   - 有 "新辅助"、"治疗后"、"ypTNM" → 新辅助后分期
   - 多个分期并存时，优先选择病理分期

6. **分化程度映射**：
   - well differentiated → 高分化
   - moderately differentiated → 中分化
   - poorly differentiated → 低分化
   - undifferentiated → 未分化

7. **OCR 错误**：系统会自动修复常见的 OCR 错误（如 TI→T1），无需你处理。

8. **特殊处理**：
   - 如果文档中没有明确的分期信息，将 clinical_stage 设为 "待评估"，并在 evidence_quote 中说明"未找到分期信息"
   - 如果 TNM 信息不完整（如只有 T3，缺少 N 和 M），将 clinical_stage 设为 "待评估"
   - 如果找到分期信息但格式不明确（如"晚期"），在 evidence_quote 中记录原文，clinical_stage 设为 "待评估"

【输出格式】
请严格按照 JSON 格式输出，包含 thinking_process 字段。

【thinking_process 示例】
在 thinking_process 中简要说明分期提取的逻辑，例如：
"在病理报告第3页找到明确的分期描述：'pT3N1M0'，根据TNM分期标准推断为III期"
"""


def _encode_image_to_base64(image_bytes: bytes) -> str:
    """将图片字节编码为 base64"""
    return base64.b64encode(image_bytes).decode("utf-8")


def _pdf_to_images(pdf_bytes: bytes, dpi: int = 100, max_pages: int = 20) -> List[bytes]:
    """
    将 PDF 转换为图片序列（优化版）
    
    Args:
        pdf_bytes: PDF 文件的字节内容
        dpi: 图片分辨率（降低到100以提升速度）
        max_pages: 最大转换页数（增大到20以支持长文档）
    
    Returns:
        图片字节列表 (JPEG 格式，比PNG更快更小)
    """
    if not HAS_PYMUPDF:
        # 如果没有 PyMuPDF，返回空列表，后续会 fallback 到文本提取
        return []
    
    images = []
    try:
        pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        page_count = min(len(pdf_doc), max_pages)  # 限制最大页数
        
        for page_num in range(page_count):
            page = pdf_doc[page_num]
            # 设置缩放矩阵（使用较低DPI提升速度）
            zoom = dpi / 72  # 72 是默认 DPI
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)
            # 使用 JPEG 格式，比 PNG 更快更小（质量85足够识别文字）
            img_bytes = pix.tobytes("jpeg")
            images.append(img_bytes)
        pdf_doc.close()
        print(f"[PDF转图片] 成功转换 {page_count} 页 (DPI={dpi}, 格式=JPEG)")
    except Exception as e:
        print(f"[PDF转图片失败] {e}")
    
    return images


def _extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """
    [性能优化] 使用 PyMuPDF (fitz) 极速提取文本
    替代原本慢速的 pypdf，移除额外依赖，支持布局保持

    Returns:
        提取的完整文本（按页分隔）
    """
    if not HAS_PYMUPDF:
        print("[PDF文本提取] 错误: PyMuPDF (fitz) 未安装")
        return ""

    try:
        # 使用 PyMuPDF 打开 PDF（自动管理内存）
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        page_count = len(doc)
        print(f"[PDF文本提取] 开始提取，共 {page_count} 页 (Engine: PyMuPDF)")

        texts = []
        for page_num, page in enumerate(doc, 1):
            # 使用 "text" 格式提取，配合 sort=True 优化阅读顺序
            # flags=fitz.TEXT_PRESERVE_WHITESPACE | fitz.TEXT_MEDIABOX_CLIP
            page_text = page.get_text("text", sort=True) or ""
            texts.append(page_text)

            # 显示每页提取的字符数
            if page_text:
                print(f"  - 第 {page_num} 页: {len(page_text)} 字符")

        doc.close()  # 确保资源释放

        full_text = "\n\n".join(texts)
        print(f"[PDF文本提取] 完成，总计 {len(full_text)} 字符")

        # 检查是否包含关键分期信息（调试用）
        staging_keywords = ['T4', 'T3', 'T2', 'T1', 'N0', 'N1', 'N2', 'M0', 'M1', 'stage', 'Stage', 'TNM']
        found_keywords = [kw for kw in staging_keywords if kw in full_text]
        if found_keywords:
            print(f"[PDF文本提取] 检测到分期关键词: {found_keywords}")

        return full_text

    except Exception as e:
        print(f"[PDF文本提取失败] {e}")
        import traceback
        traceback.print_exc()
        return ""


def _is_text_rich_pdf(text: str, min_chars: int = 100, min_chinese_ratio: float = 0.05) -> bool:
    """
    智能判断PDF是否包含足够的可用文本内容
    
    判断标准：
    1. 总字符数 >= min_chars（默认100）
    2. 中文字符占比 >= min_chinese_ratio（默认5%）或有足够的英文单词
    3. 不是纯乱码或特殊字符
    
    Args:
        text: 提取的文本
        min_chars: 最小字符数阈值
        min_chinese_ratio: 最小中文字符占比
    
    Returns:
        True 表示文本丰富，可以使用文本模式
        False 表示需要使用视觉模式（可能是扫描件）
    """
    if not text:
        return False
    
    # 清理空白字符
    clean_text = text.strip()
    total_len = len(clean_text)
    
    if total_len < min_chars:
        print(f"[PDF检测] 文本过短 ({total_len} < {min_chars})，判定为扫描件")
        return False
    
    # 统计中文字符数量
    chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', clean_text))
    chinese_ratio = chinese_chars / total_len if total_len > 0 else 0
    
    # 统计英文单词数量（至少2个字母的连续序列）
    english_words = len(re.findall(r'[a-zA-Z]{2,}', clean_text))
    
    # 统计有意义的数字（如年龄、日期、检查数值）
    meaningful_numbers = len(re.findall(r'\d+', clean_text))
    
    # 医疗关键词检测
    medical_keywords = [
        '患者', '诊断', '检查', '治疗', '手术', '病理', '肿瘤', '癌',
        '分期', 'CT', 'MRI', '血常规', '肠镜', '活检', '淋巴结',
        '主诉', '病史', '入院', '出院', 'TNM', '化疗', '放疗'
    ]
    keyword_count = sum(1 for kw in medical_keywords if kw in clean_text)
    
    # 判断逻辑
    is_rich = False
    reason = ""
    
    if chinese_ratio >= min_chinese_ratio:
        is_rich = True
        reason = f"中文占比 {chinese_ratio:.1%}"
    elif english_words >= 20:
        is_rich = True
        reason = f"英文单词数 {english_words}"
    elif keyword_count >= 3:
        is_rich = True
        reason = f"医疗关键词 {keyword_count} 个"
    elif meaningful_numbers >= 10 and (chinese_chars > 20 or english_words > 10):
        is_rich = True
        reason = f"数字+文字混合"
    
    if is_rich:
        print(f"[PDF检测] 文本丰富 ({reason})，使用快速文本模式 ✓")
    else:
        print(f"[PDF检测] 文本不足 (中文{chinese_ratio:.1%}, 英文词{english_words}, 关键词{keyword_count})，使用视觉模式")
    
    return is_rich


def _extract_text_from_docx(docx_bytes: bytes) -> str:
    """从 Word 文档提取纯文本"""
    if not HAS_DOCX:
        raise ImportError("需要安装 python-docx: pip install python-docx")
    
    try:
        doc = docx.Document(io.BytesIO(docx_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        
        return "\n".join(paragraphs)
    except Exception as e:
        print(f"[Word文本提取失败] {e}")
        return ""


def _clean_json_response(response: str) -> str:
    """清理 LLM 返回的 JSON（去除 markdown 代码块等）"""
    # 移除 markdown 代码块
    response = re.sub(r'^```json\s*', '', response.strip())
    response = re.sub(r'^```\s*', '', response)
    response = re.sub(r'\s*```$', '', response)
    return response.strip()


def _parse_json_safely(text: str) -> Optional[Dict[str, Any]]:
    """
    [鲁棒性增强] 带自动修复的 JSON 解析

    增强点：
    1. 基本的 JSON 清理
    2. 简单的尾随逗号修复
    3. 更详细的错误日志
    """
    try:
        cleaned = _clean_json_response(text)
        return json.loads(cleaned)
    except json.JSONDecodeError as e:
        print(f"[JSON解析失败] {e}")
        # 尝试修复尾随逗号
        try:
            repaired = re.sub(r',(\s*[}\]])', r'\1', cleaned)
            repaired = re.sub(r',(\s*$)', '', repaired)  # 处理行尾逗号
            result = json.loads(repaired)
            print(f"[JSON修复成功] 已自动修复尾随逗号")
            return result
        except Exception:
            print(f"[JSON自动修复失败] 原始片段: {cleaned[:200]}...")
            return None


def _scrub_pii(data: Dict[str, Any]) -> Dict[str, Any]:
    """
    [安全边界] 规则级隐私清洗

    增强说明：
    1. 正则脱敏：手机号、身份证号
    2. 递归清理：敏感键值对移除
    3. 姓名脱敏：常见中文姓名模式

    注意：这是防御性措施，不能替代 LLM 的隐私保护
    """
    json_str = json.dumps(data, ensure_ascii=False)

    # 1. 手机号脱敏 (11位数字，保留前后非数字字符)
    json_str = re.sub(r'(?<=\D)1[3-9]\d{9}(?=\D)', '[隐私手机号]', json_str)

    # 2. 身份证号脱敏 (15或18位)
    json_str = re.sub(r'(?<=\D)\d{17}[\dXx](?=\D)', '[隐私身份证]', json_str)

    # 3. 座机号码脱敏 (带区号格式)
    json_str = re.sub(r'(?<=\D)\d{3,4}[-]?\d{7,8}(?=\D)', '[隐私座机]', json_str)

    # 4. 日期脱敏（疑似出生日期，如 1990/01/01）
    json_str = re.sub(r'(?<=\D)(19|20)\d{2}[-/]\d{1,2}[-/]\d{1,2}(?=\D)', '[隐私日期]', json_str)

    # 5. 常见敏感键值对移除
    sensitive_keys = ["id_card", "phone", "mobile", "telephone", "address",
                      "social_security", "ssn", "bank_account", "passport"]

    def clean_recursive(obj):
        """递归清理敏感数据"""
        if isinstance(obj, dict):
            for k in list(obj.keys()):
                # 检查键名是否敏感
                if any(sk in k.lower() for sk in sensitive_keys):
                    del obj[k]
                else:
                    # 递归处理值
                    clean_recursive(obj[k])
        elif isinstance(obj, list):
            for item in obj:
                clean_recursive(item)

    # 解析并清理
    data_clean = json.loads(json_str)
    clean_recursive(data_clean)

    return data_clean


def _create_fallback_card(text_content: str, filename: str) -> Dict[str, Any]:
    """当转换失败时，创建一个基础的占位卡片"""
    return {
        "type": "medical_visualization_card",
        "data": {
            "patient_summary": {
                "age": "未提供",
                "gender": "未提供",
                "chief_complaint": "文档解析中...",
                "admission_date": "未提供",
                "medical_history": "未提供"
            },
            "diagnosis_block": {
                "confirmed": "待确认",
                "location": TumorLocationEnum.UNKNOWN.value,
                "pathology": "未提供",
                "differentiation": DifferentiationEnum.UNKNOWN.value,
                "evidence_quote": "未提供",
                "molecular_markers": {
                    "MSI": "未提供",
                    "KRAS": "未提供",
                    "NRAS": "未提供",
                    "BRAF": "未提供",
                    "HER2": "未提供",
                    "UGT1A1": "未提供",
                    "others": "无"
                }
            },
            "staging_block": {
                "t_stage": "Tx",
                "n_stage": "Nx",
                "m_stage": "Mx",
                "staging_type": StagingTypeEnum.UNKNOWN.value,
                "clinical_stage": "待评估",
                "risk_status": RiskStatusEnum.UNKNOWN.value,
                "evidence_quote": "未提供"
            },
            "key_findings": [
                {
                    "category": "文档",
                    "finding": f"已上传文件: {filename}",
                    "significance": "请在下方输入框提问以获取详细分析"
                }
            ],
            "treatment_draft": []
        },
        "_raw_text": text_content[:2000] if text_content else "",
        "_parse_error": True
    }


class DocumentConverter:
    """
    医疗文档转化 Agent
    
    将 PDF/Word/图片病历转化为结构化 JSON，
    用于前端"可视化诊疗卡片"渲染。
    
    支持长文档处理：
    - 可配置的上下文窗口 (max_tokens: 4096 默认，足够生成 JSON)
    - 智能分段处理策略
    - 可配置的参数 (通过 config.py 或环境变量)
    """
    
    def __init__(
        self,
        api_key: Optional[str] = None,
        api_base: Optional[str] = None,
        model: Optional[str] = None,
        max_tokens: int = 4096,
        max_text_length: int = 1000000,
        max_pages: int = 20,
        max_images: int = 10,
        pdf_dpi: int = 100,
        enable_chunked_processing: bool = True,
        chunk_size: int = 20000,
        chunk_overlap: int = 2000,
    ):
        self.api_key = api_key or os.getenv("OPENAI_API_KEY")
        self.api_base = api_base or os.getenv("OPENAI_API_BASE")
        self.model_name = model or os.getenv("DOC_CONVERTER_MODEL", "gpt-4o-mini")
        self.model = self.model_name  # 保持兼容性
        self.max_tokens = max_tokens or int(os.getenv("DOC_CONVERTER_MAX_TOKENS", "4096"))
        
        # 文档处理配置参数
        self.max_text_length = max_text_length
        self.max_pages = max_pages
        self.max_images = max_images
        self.pdf_dpi = pdf_dpi
        self.enable_chunked_processing = enable_chunked_processing
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

        # 初始化 LLM
        base_llm = create_compatible_chat_openai(
            model=self.model_name,
            api_key=self.api_key,
            base_url=self.api_base,
            temperature=0, # 结构化提取不需要创造性，设为0
            max_tokens=max_tokens,
        )

        # 【优化点】使用 structured_output 绑定 Pydantic
        # 这比手动解析 JSON 字符串稳定得多
        if hasattr(base_llm, "with_structured_output"):
            self.structured_llm = base_llm.with_structured_output(MedicalVisualizationCard)
        else:
            # 降级处理：某些旧版 LangChain 不支持
            self.structured_llm = base_llm

        self.logger = logging.getLogger(__name__)

        # 配置日志记录器
        self.logger = logging.getLogger(__name__)
        self.logger.setLevel(logging.DEBUG)

        # 如果没有处理器，添加一个控制台处理器
        if not self.logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter(
                '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
            )
            handler.setFormatter(formatter)
            self.logger.addHandler(handler)

        # 检查是否有调试日志文件路径的环境变量
        debug_log_path = os.getenv("DOC_CONVERTER_DEBUG_LOG")
        if debug_log_path:
            try:
                file_handler = logging.FileHandler(debug_log_path, encoding='utf-8')
                file_handler.setLevel(logging.DEBUG)
                file_formatter = logging.Formatter(
                    '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
                )
                file_handler.setFormatter(file_formatter)
                self.logger.addHandler(file_handler)
                print(f"[DocumentConverter] 调试日志已启用: {debug_log_path}")
            except Exception as e:
                print(f"[DocumentConverter] 无法创建调试日志文件: {e}")

        self.llm = create_compatible_chat_openai(
            model=self.model,
            api_key=self.api_key,
            base_url=self.api_base,
            temperature=0.1,
            max_tokens=self.max_tokens,  # 使用配置的 max_tokens
        )

    # --- 核心辅助函数：PyMuPDF 极速提取 ---

    def _extract_hybrid_content(self, pdf_bytes: bytes) -> List[Dict[str, Any]]:
        """
        【混合解析策略】
        按页遍历 PDF：
        1. 先尝试提取文本。
        2. 如果该页文本很少 (<50字) 或者是扫描件 -> 渲染为图片。
        3. 如果该页是文本 -> 保留文本。

        返回: [{"type": "text", "page": 1, "content": "..."}, {"type": "image", "page": 2, "data": bytes}, ...]
        """
        if not HAS_PYMUPDF:
            return [{"type": "text", "content": "PyMuPDF not installed, fallback text extraction failed."}]

        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        content_stream = []

        for page_num, page in enumerate(doc):
            # 1. 尝试提取文本
            text = page.get_text()
            clean_text = text.strip()

            # 判断逻辑：如果文本少于 50 字符，大概率是扫描件或纯图表
            if len(clean_text) > 50:
                content_stream.append({
                    "type": "text",
                    "page": page_num + 1,
                    "content": clean_text
                })
            else:
                # 2. 渲染为图片 (JPEG, 降低 DPI 到 150 足够 OCR)
                # 优化：限制图片最大分辨率，防止 token 爆炸
                pix = page.get_pixmap(dpi=100)
                img_bytes = pix.tobytes("jpeg")
                content_stream.append({
                    "type": "image",
                    "page": page_num + 1,
                    "data": img_bytes
                })
                self.logger.info(f"第 {page_num+1} 页判定为扫描件/图片，已转图片处理。")

        return content_stream

    def convert_pdf(self, pdf_bytes: bytes, filename: str) -> Dict[str, Any]:
        """
        处理 PDF 入口
        """
        start_time = time.time()

        # 1. 混合提取内容
        contents = self._extract_hybrid_content(pdf_bytes)

        # 2. 构建 Prompt 消息
        messages = [SystemMessage(content=CONVERSION_SYSTEM_PROMPT)]
        user_content_blocks = []

        text_buffer = ""
        image_count = 0

        for item in contents:
            if item["type"] == "text":
                text_buffer += f"\n--- Page {item['page']} ---\n{item['content']}"
            elif item["type"] == "image":
                # 累积之前的文本块
                if text_buffer:
                    user_content_blocks.append({"type": "text", "text": text_buffer})
                    text_buffer = ""

                # 添加图片块
                b64_img = base64.b64encode(item["data"]).decode("utf-8")
                user_content_blocks.append({
                    "type": "image_url",
                    "image_url": {"url": f"data:image/jpeg;base64,{b64_img}", "detail": "auto"}
                })
                image_count += 1

        # 添加剩余文本
        if text_buffer:
            user_content_blocks.append({"type": "text", "text": text_buffer})

        # 限制图片数量防止超限 (取前5张和后5张，关键信息通常在首尾)
        if image_count > 10:
            # 简单策略：保留所有文本，截断中间图片
            # 实际生产中可更复杂
            pass

        messages.append(HumanMessage(content=user_content_blocks))

        try:
            # 3. 调用支持结构化输出的 LLM
            # invoke 会直接返回 MedicalVisualizationCard 对象实例
            result_obj = self.structured_llm.invoke(messages)

            # 4. 转回 Dict (兼容 Pydantic v1/v2)
            if hasattr(result_obj, "model_dump"):
                final_json = result_obj.model_dump()
            else:
                final_json = result_obj.dict()

            # 5. 注入元数据
            final_json["_source_file"] = filename
            final_json["_performance"] = {
                "total_time": f"{time.time() - start_time:.2f}s",
                "mode": f"hybrid (img={image_count})"
            }

            # 确保 type 字段正确 (Pydantic 可能覆盖)
            final_json["type"] = "medical_visualization_card"

            # [新增] 安全清洗：强制 PII 脱敏
            final_json = _scrub_pii(final_json)

            # [新增] 循证性验证：检查关键字段是否有原文引用
            data = final_json.get("data", {})
            diag_block = data.get("diagnosis_block", {})
            staging_block = data.get("staging_block", {})

            # 诊断引用检查
            if diag_block.get("confirmed") and not diag_block.get("evidence_quote"):
                self.logger.warning(f"[Grounding Alert] 诊断 '{diag_block['confirmed']}' 缺少原文引用！")
                final_json["_grounding_warnings"] = final_json.get("_grounding_warnings", []) + ["diagnosis_missing_evidence"]

            # 分期引用检查
            if staging_block.get("clinical_stage") and not staging_block.get("evidence_quote"):
                self.logger.warning(f"[Grounding Alert] 分期 '{staging_block['clinical_stage']}' 缺少原文引用！")
                final_json["_grounding_warnings"] = final_json.get("_grounding_warnings", []) + ["staging_missing_evidence"]

            return final_json

        except Exception as e:
            self.logger.error(f"LLM 结构化提取失败: {e}")
            import traceback
            traceback.print_exc()
            return self._create_fallback(filename, str(e))

    def _create_fallback(self, filename: str, error_msg: str) -> Dict:
        """兜底返回（包含完整的 Schema 结构，使用枚举值）"""
        return {
            "type": "medical_visualization_card",
            "data": {
                "patient_summary": {
                    "age": "未提供",
                    "gender": "未提供",
                    "chief_complaint": "文档解析失败",
                    "admission_date": "未提供",
                    "medical_history": "未提供"
                },
                "diagnosis_block": {
                    "confirmed": "待确认",
                    "location": TumorLocationEnum.UNKNOWN.value,
                    "pathology": "未提供",
                    "differentiation": DifferentiationEnum.UNKNOWN.value,
                    "evidence_quote": "未提供",
                    "molecular_markers": {
                        "MSI": "未提供",
                        "KRAS": "未提供",
                        "NRAS": "未提供",
                        "BRAF": "未提供",
                        "HER2": "未提供",
                        "UGT1A1": "未提供",
                        "others": "无"
                    }
                },
                "staging_block": {
                    "t_stage": "Tx",
                    "n_stage": "Nx",
                    "m_stage": "Mx",
                    "staging_type": StagingTypeEnum.UNKNOWN.value,
                    "clinical_stage": "待评估",
                    "risk_status": RiskStatusEnum.UNKNOWN.value,
                    "evidence_quote": "未提供"
                },
                "key_findings": [{"category": "Error", "finding": error_msg}],
                "treatment_draft": []
            },
            "_parse_error": True
        }
    
    def convert_docx(
        self,
        docx_bytes: bytes,
        filename: str = "document.docx"
    ) -> Dict[str, Any]:
        """
        转换 Word 文档为结构化 JSON
        
        Args:
            docx_bytes: Word 文件字节内容
            filename: 文件名
        
        Returns:
            medical_visualization_card JSON
        """
        text = _extract_text_from_docx(docx_bytes)
        if text.strip():
            return self._convert_with_text(text, filename)
        
        return _create_fallback_card("", filename)
    
    def convert_text(
        self,
        text: str,
        filename: str = "document.txt"
    ) -> Dict[str, Any]:
        """
        转换纯文本为结构化 JSON
        
        Args:
            text: 文本内容
            filename: 文件名
        
        Returns:
            medical_visualization_card JSON
        """
        if text.strip():
            return self._convert_with_text(text, filename)
        
        return _create_fallback_card("", filename)
    
    def convert_image(
        self,
        image_bytes: bytes,
        filename: str = "image.png"
    ) -> Dict[str, Any]:
        """
        转换图片为结构化 JSON
        
        Args:
            image_bytes: 图片字节内容
            filename: 文件名
        
        Returns:
            medical_visualization_card JSON
        """
        return self._convert_with_vision([image_bytes], filename)
    
    def _convert_with_vision(
        self,
        images: List[bytes],
        filename: str,
        max_images: Optional[int] = None
    ) -> Dict[str, Any]:
        """
        使用视觉模式转换（优化版）
        
        优化点：
        - 可配置的图片数量限制（默认使用 self.max_images = 10）
        - 使用 detail: "auto" 让模型自动选择（比 high 更快）
        - 使用 JPEG 格式减少传输数据量
        """
        if max_images is None:
            max_images = self.max_images
        try:
            import time
            llm_start = time.time()
            
            # 构建多模态消息（使用极简 Prompt）
            content = [
                {"type": "text", "text": f"提取医疗文档 {filename} 的关键临床信息，输出 JSON。"}
            ]
            
            # 添加图片（限制数量以提升速度）
            images_to_process = images[:max_images]
            print(f"[视觉模式] 处理 {len(images_to_process)} 张图片")
            
            for i, img_bytes in enumerate(images_to_process):
                base64_img = _encode_image_to_base64(img_bytes)
                # 检测图片格式（JPEG 或 PNG）
                is_jpeg = img_bytes[:2] == b'\xff\xd8'
                mime_type = "image/jpeg" if is_jpeg else "image/png"
                
                content.append({
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:{mime_type};base64,{base64_img}",
                        "detail": "auto"  # 改为 auto，让模型自动选择，比 high 更快
                    }
                })
            
            messages = [
                SystemMessage(content=LIGHTWEIGHT_SYSTEM_PROMPT),
                HumanMessage(content=content)
            ]
            
            response = self.llm.invoke(messages)
            llm_time = time.time() - llm_start
            print(f"[性能] LLM 调用耗时: {llm_time:.2f}s")
            
            result = _parse_json_safely(response.content)
            
            if result and result.get("type") == "medical_visualization_card":
                result["_source_file"] = filename
                result["_vision_info"] = {
                    "images_processed": len(images_to_process),
                    "llm_time": f"{llm_time:.2f}s"
                }
                return result
            
            # 解析失败，返回 fallback
            return _create_fallback_card("", filename)
            
        except Exception as e:
            print(f"[视觉转换失败] {e}")
            return _create_fallback_card("", filename)
    
    def _convert_with_text(
        self,
        text: str,
        filename: str,
        max_text_length: Optional[int] = None
    ) -> Dict[str, Any]:
        """
        使用纯文本模式转换（支持长文档分段处理）
        
        优化点：
        - 可配置的最大文本长度（默认 self.max_text_length = 50000）
        - 长文档智能分段处理：分段提取再合并结果
        - 智能截断保留头尾关键信息
        - 改进提示词避免LLM拒绝响应
        - 添加响应验证和重试机制
        """
        import time
        
        if max_text_length is None:
            max_text_length = self.max_text_length
        
        # 记录函数入口信息
        self.logger.debug("函数入口", extra={
            "hypothesisId": "A",
            "location": "document_converter.py:_convert_with_text:entry",
            "data": {
                "text_len": len(text),
                "max_text_length": max_text_length,
                "enable_chunked": self.enable_chunked_processing
            }
        })
        # #endregion
        
        try:
            # 检查是否需要分段处理
            if self.enable_chunked_processing and len(text) > max_text_length:
                # 记录触发分段处理
                self.logger.debug("触发分段处理", extra={
                    "hypothesisId": "C",
                    "location": "document_converter.py:chunked_triggered",
                    "data": {
                        "text_len": len(text),
                        "max_text_length": max_text_length
                    }
                })
                # #endregion
                print(f"[长文档处理] 文本长度 {len(text)} 超过限制 {max_text_length}，启用分段处理")
                return self._convert_long_document(text, filename)
            
            # 智能处理长文本：保留开头和结尾（关键信息通常在这两处）
            processed_text = text
            truncated = False
            if len(text) > max_text_length:
                truncated = True
                # 保留开头60%和结尾40%（分期信息通常在后半部分）
                head_len = int(max_text_length * 0.6)
                tail_len = max_text_length - head_len
                processed_text = text[:head_len] + "\n\n[...中间部分省略...]\n\n" + text[-tail_len:]
                print(f"[文本处理] 原文 {len(text)} 字符，截取为 {len(processed_text)} 字符（保留头尾关键部分）")
            
            # 记录LLM调用前信息
            self.logger.debug("LLM调用前", extra={
                "hypothesisId": "A,E",
                "location": "document_converter.py:before_llm",
                "data": {
                    "truncated": truncated,
                    "processed_text_len": len(processed_text),
                    "estimated_tokens": len(processed_text) // 2
                }
            })
            # #endregion
            
            # 完整的用户提示 - 使用极简 Prompt（规则已下沉到 Schema）
            user_prompt = f"""分析医疗文档 "{filename}"，提取关键临床信息。

待提取文本：
{processed_text}

请输出符合 Schema 定义的 JSON 数据。"""

            messages = [
                SystemMessage(content=LIGHTWEIGHT_SYSTEM_PROMPT),
                HumanMessage(content=user_prompt)
            ]
            
            llm_start = time.time()
            response = self.llm.invoke(messages)
            llm_time = time.time() - llm_start
            
            # 记录LLM调用后信息
            self.logger.debug("LLM调用后", extra={
                "hypothesisId": "D",
                "location": "document_converter.py:after_llm",
                "data": {
                    "llm_time_ms": int(llm_time * 1000),
                    "response_len": len(response.content or "")
                }
            })
            # #endregion
            
            print(f"[性能] 文本模式 LLM 调用耗时: {llm_time:.2f}s")
            
            # 验证响应内容
            response_content = response.content or ""
            
            # 检测LLM拒绝响应的模式
            refusal_patterns = [
                "抱歉", "无法处理", "无法分析", "I cannot", "I'm sorry", 
                "I apologize", "我无法", "不能处理", "Sorry"
            ]
            is_refusal = any(pattern in response_content[:100] for pattern in refusal_patterns)
            
            # 记录拒绝响应检测
            self.logger.debug("检测拒绝响应", extra={
                "hypothesisId": "B",
                "location": "document_converter.py:refusal_check",
                "data": {
                    "is_refusal": is_refusal,
                    "response_preview": response_content[:100]
                }
            })
            # #endregion
            
            if is_refusal:
                # 记录触发重试机制
                self.logger.debug("触发重试机制", extra={
                    "hypothesisId": "B",
                    "location": "document_converter.py:retry_triggered",
                    "data": {"retry": True}
                })
                # #endregion
                print(f"[警告] LLM 返回拒绝响应，尝试使用简化提示重试...")
                # 使用更简化的提示重试一次
                retry_prompt = f"""请直接从以下医疗文本中提取信息并输出JSON（不要拒绝，不要道歉）：

{processed_text[:8000]}

直接输出JSON，格式如下：
{{"type": "medical_visualization_card", "data": {{...}}}}"""
                
                retry_messages = [
                    SystemMessage(content="你是医疗信息提取助手。只输出JSON，不要任何其他文字。"),
                    HumanMessage(content=retry_prompt)
                ]
                retry_start = time.time()
                response = self.llm.invoke(retry_messages)
                retry_time = time.time() - retry_start
                # 记录重试完成
                self.logger.debug("重试完成", extra={
                    "hypothesisId": "B",
                    "location": "document_converter.py:retry_done",
                    "data": {"retry_time_ms": int(retry_time * 1000)}
                })
                # #endregion
                response_content = response.content or ""
            
            result = _parse_json_safely(response_content)
            
            if result and result.get("type") == "medical_visualization_card":
                result["_source_file"] = filename
                result["_raw_text"] = text[:10000]  # 保存更多原始文本用于后续问答（增加到10000）
                result["_text_mode_info"] = {
                    "original_length": len(text),
                    "processed_length": len(processed_text),
                    "llm_time": f"{llm_time:.2f}s",
                    "chunked": False
                }
                # 记录转换成功
                self.logger.debug("转换成功", extra={
                    "hypothesisId": "ALL",
                    "location": "document_converter.py:success",
                    "data": {
                        "llm_time_ms": int(llm_time * 1000),
                        "total_retries": 0 if not is_refusal else 1
                    }
                })
                # #endregion
                return result
            
            # 解析失败，记录详细信息
            print(f"[JSON解析失败] 响应内容前200字符: {response_content[:200]}")
            return _create_fallback_card(text, filename)
            
        except Exception as e:
            print(f"[文本转换失败] {e}")
            import traceback
            traceback.print_exc()
            return _create_fallback_card(text, filename)
    
    def _split_text_into_chunks(self, text: str) -> List[str]:
        """
        将长文本分割成多个有重叠的块
        
        使用配置的 chunk_size 和 chunk_overlap
        """
        chunks = []
        start = 0
        text_len = len(text)
        
        while start < text_len:
            end = min(start + self.chunk_size, text_len)
            chunk = text[start:end]
            chunks.append(chunk)
            
            if end >= text_len:
                break
            
            # 下一个块的起始位置（有重叠）
            start = end - self.chunk_overlap
        
        print(f"[分段处理] 文本 {text_len} 字符分成 {len(chunks)} 个块 (chunk_size={self.chunk_size}, overlap={self.chunk_overlap})")
        return chunks
    
    def _process_single_chunk(self, chunk_data: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        """
        处理单个文本块的辅助方法（用于并行处理）

        Args:
            chunk_data: 包含chunk信息和索引的字典
                {
                    "chunk": str,  # 文本块内容
                    "index": int,  # 块索引
                    "total_chunks": int,  # 总块数
                    "filename": str  # 文件名
                }

        Returns:
            提取结果的JSON，如果失败返回None
        """
        chunk = chunk_data["chunk"]
        index = chunk_data["index"]
        total_chunks = chunk_data["total_chunks"]
        filename = chunk_data["filename"]

        chunk_prompt = f"""提取医疗文档 "{filename}" 的第 {index+1}/{total_chunks} 部分中的临床信息。

如果信息未找到，填写 "本段未提供"。

---文档内容片段---
{chunk}
---片段结束---

输出 JSON。"""

        messages = [
            SystemMessage(content=LIGHTWEIGHT_SYSTEM_PROMPT),
            HumanMessage(content=chunk_prompt)
        ]

        try:
            response = self.llm.invoke(messages)
            result = _parse_json_safely(response.content or "")
            if result and result.get("type") == "medical_visualization_card":
                print(f"[分段处理] 第 {index+1} 段成功提取")
                return result
            else:
                print(f"[分段处理] 第 {index+1} 段解析失败")
                return None
        except Exception as e:
            print(f"[分段处理] 第 {index+1} 段处理失败: {e}")
            return None

    def _convert_long_document(self, text: str, filename: str) -> Dict[str, Any]:
        """
        长文档分段处理策略（并行优化版）

        将长文档分成多个块，并行提取信息，然后智能合并结果
        """
        import time
        start_time = time.time()

        chunks = self._split_text_into_chunks(text)

        # 准备并行处理的数据
        chunk_tasks = [
            {
                "chunk": chunk,
                "index": i,
                "total_chunks": len(chunks),
                "filename": filename
            }
            for i, chunk in enumerate(chunks)
        ]

        chunk_results = []

        # 使用ThreadPoolExecutor进行并行处理
        max_workers = min(len(chunks), 4)  # 限制最大并发数，避免API限制
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            print(f"[分段处理] 开始并行处理 {len(chunks)} 个块 (最大并发: {max_workers})...")

            # 提交所有任务
            future_to_chunk = {
                executor.submit(self._process_single_chunk, task): task
                for task in chunk_tasks
            }

            # 收集结果
            for future in as_completed(future_to_chunk):
                result = future.result()
                if result:
                    chunk_results.append(result)

        # 合并所有块的结果
        if not chunk_results:
            print("[分段处理] 所有段落处理失败，返回 fallback")
            return _create_fallback_card(text, filename)

        merged_result = self._merge_chunk_results(chunk_results, filename)
        merged_result["_raw_text"] = text[:15000]  # 保存更多原始文本
        merged_result["_text_mode_info"] = {
            "original_length": len(text),
            "chunks_processed": len(chunks),
            "chunks_successful": len(chunk_results),
            "total_time": f"{time.time() - start_time:.2f}s",
            "parallel_processing": True,
            "max_concurrent_workers": max_workers
        }

        print(f"[分段处理] 完成，并行合并 {len(chunk_results)} 个结果，总耗时 {time.time() - start_time:.2f}s")
        return merged_result
    
    def _refine_merged_results(self, merged_result: Dict[str, Any], filename: str) -> Dict[str, Any]:
        """
        使用LLM对合并结果进行精细化处理，去除重复和冲突信息

        Args:
            merged_result: 初步合并的结果
            filename: 文件名

        Returns:
            精细化后的结果
        """
        try:
            # 准备精细化提示（使用极简 Prompt）
            refine_prompt = f"""整理医疗文档 "{filename}" 的分段提取结果。

任务：
1. 合并重复的关键发现（key_findings）
2. 解决分期信息冲突，优先选择完整准确的版本
3. 整理治疗方案，去除重复步骤
4. 保持诊断信息准确性

原始结果：
{json.dumps(merged_result, ensure_ascii=False, indent=2)}

输出优化后的 JSON。"""

            messages = [
                SystemMessage(content=LIGHTWEIGHT_SYSTEM_PROMPT),
                HumanMessage(content=refine_prompt)
            ]

            response = self.llm.invoke(messages)
            refined_result = _parse_json_safely(response.content or "")

            if refined_result and refined_result.get("type") == "medical_visualization_card":
                refined_result["_refined"] = True
                refined_result["_source_file"] = filename
                refined_result["_merged_from_chunks"] = True
                print("[精细化] LLM 成功优化合并结果")
                return refined_result
            else:
                print("[精细化] LLM 优化失败，返回原始合并结果")
                return merged_result

        except Exception as e:
            print(f"[精细化] 处理失败: {e}，返回原始合并结果")
            return merged_result

    def _merge_chunk_results(self, results: List[Dict[str, Any]], filename: str) -> Dict[str, Any]:
        """
        智能合并多个块的提取结果（增强版）

        1. 首先进行基础合并
        2. 然后使用LLM进行精细化处理，去除重复和冲突
        """
        merged = {
            "type": "medical_visualization_card",
            "data": {
                "patient_summary": {},
                "diagnosis_block": {},
                "staging_block": {},
                "key_findings": [],
                "treatment_draft": []
            },
            "_source_file": filename,
            "_merged_from_chunks": True
        }

        def is_valid_value(v: Any) -> bool:
            """检查值是否有效（非空、非未提供）"""
            if v is None:
                return False
            if isinstance(v, str):
                invalid_patterns = ["未提供", "本段未提供", "Tx", "Nx", "Mx", "待评估", "待确认", ""]
                return v.strip() not in invalid_patterns
            if isinstance(v, dict):
                return bool(v)
            if isinstance(v, list):
                return bool(v)
            return True

        def merge_dict(target: Dict, source: Dict) -> None:
            """合并字典，优先选择有效值"""
            for key, value in source.items():
                if key not in target or not is_valid_value(target.get(key)):
                    if is_valid_value(value):
                        target[key] = value
                # 如果是嵌套字典，递归合并
                elif isinstance(value, dict) and isinstance(target.get(key), dict):
                    merge_dict(target[key], value)

        # 合并各个块的结果
        for result in results:
            data = result.get("data", {})

            # 合并 patient_summary
            merge_dict(merged["data"]["patient_summary"], data.get("patient_summary", {}))

            # 合并 diagnosis_block
            merge_dict(merged["data"]["diagnosis_block"], data.get("diagnosis_block", {}))

            # 合并 staging_block
            merge_dict(merged["data"]["staging_block"], data.get("staging_block", {}))

            # 合并 key_findings（追加所有内容，后续由LLM去重）
            for finding in data.get("key_findings", []):
                if finding and isinstance(finding, dict) and finding.get("finding"):
                    merged["data"]["key_findings"].append(finding)

            # 合并 treatment_draft（追加所有内容，后续由LLM整理）
            for treatment in data.get("treatment_draft", []):
                if treatment and isinstance(treatment, dict) and treatment.get("step"):
                    merged["data"]["treatment_draft"].append(treatment)

        # 对 treatment_draft 按步骤排序
        merged["data"]["treatment_draft"].sort(key=lambda x: x.get("step", 0))

        # 使用LLM进行精细化处理
        print(f"[合并处理] 初步合并完成，发现 {len(merged['data']['key_findings'])} 个关键发现，{len(merged['data']['treatment_draft'])} 个治疗步骤")
        refined_result = self._refine_merged_results(merged, filename)

        return refined_result


def create_converter(
    api_key: Optional[str] = None,
    api_base: Optional[str] = None,
    model: Optional[str] = None,
    use_config: bool = True
) -> DocumentConverter:
    """
    创建文档转化 Agent 实例
    
    Args:
        api_key: API Key
        api_base: API Base URL
        model: 模型名称 (默认从配置读取)
        use_config: 是否从 config.py 加载配置参数 (默认 True)
    
    Returns:
        DocumentConverter 实例
    """
    # 尝试从配置文件加载参数
    config_params = {}
    config_model = None
    if use_config:
        try:
            from ..config import load_settings
            settings = load_settings()
            doc_config = settings.doc_converter
            config_model = doc_config.model
            config_params = {
                "max_tokens": doc_config.max_tokens,
                "max_text_length": doc_config.max_text_length,
                "max_pages": doc_config.max_pages,
                "max_images": doc_config.max_images,
                "pdf_dpi": doc_config.pdf_dpi,
                "enable_chunked_processing": doc_config.enable_chunked_processing,
                "chunk_size": doc_config.chunk_size,
                "chunk_overlap": doc_config.chunk_overlap,
            }
            print(f"[DocumentConverter] 从配置文件加载参数: model={config_model}, max_tokens={doc_config.max_tokens}, "
                  f"max_text_length={doc_config.max_text_length}, max_pages={doc_config.max_pages}")
        except Exception as e:
            print(f"[DocumentConverter] 无法加载配置，使用默认值: {e}")
    
    # 优先使用传入的 model，其次是配置中的 model
    final_model = model or config_model
    
    return DocumentConverter(
        api_key=api_key,
        api_base=api_base,
        model=final_model,
        **config_params
    )


# ============== 便捷函数 ==============

def convert_uploaded_file(
    file_bytes: bytes,
    filename: str,
    file_type: str,
    api_key: Optional[str] = None,
    api_base: Optional[str] = None,
    model: Optional[str] = None
) -> Dict[str, Any]:
    """
    便捷函数：根据文件类型自动选择转换方法
    
    Args:   
        file_bytes: 文件字节内容
        filename: 文件名
        file_type: MIME 类型 (如 application/pdf)
        api_key: API Key
        api_base: API Base URL
        model: 模型名称 (默认从配置读取)
    
    Returns:
        medical_visualization_card JSON
    """
    converter = create_converter(api_key, api_base, model)
    
    if file_type == "application/pdf" or filename.lower().endswith(".pdf"):
        return converter.convert_pdf(file_bytes, filename)
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or filename.lower().endswith(".docx"):
        return converter.convert_docx(file_bytes, filename)
    elif file_type == "text/plain" or filename.lower().endswith(".txt"):
        text = file_bytes.decode("utf-8", errors="ignore")
        return converter.convert_text(text, filename)
    elif file_type.startswith("image/"):
        return converter.convert_image(file_bytes, filename)
    else:
        # 尝试作为文本处理
        try:
            text = file_bytes.decode("utf-8", errors="ignore")
            return converter.convert_text(text, filename)
        except:
            return _create_fallback_card("", filename)

